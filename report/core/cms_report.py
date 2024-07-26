import io
import typing as t
import time
import uuid

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from mailmerge import MailMerge
from docxcompose.composer import Composer

from docx.shared import Cm
from core.base_report import BaseReport


class CmsReport(BaseReport):
    Color = {
        "注意": "ffff00",
        "告警": "ffc000",
        "故障": "ff0000"
    }

    def __init__(self, base_path, save_path, **kwargs):
        """
        :param base_path: docx模版路径
        :param save_path: 生成docx保存路径
        """
        self.base_template_path = base_path
        self.tmp_save_path = save_path
        self.document = None
        self.kwargs = kwargs
        self.merge_data = kwargs.get("merge_data") or self.get_merge_kwargs()
        self.turbine_data = kwargs.get("turbine_data") or self.get_turbine_kwargs()

    def get_merge_kwargs(self) -> t.Dict[str, t.Any]:
        kwargs = {
            "farmName": "测试风场",
            "author": "隔壁老王",
            "checker": "隔壁老张",
            "approver": "隔壁老曾",
            "createTime": "2024-01-24",
            "fanNum": "35",
            "fanModel": "xxx xxx/xxv1",
            "startTime": "2023-12-01",
            "endTime": "2023-12-30",
            "gearbox": "水平布局"
        }
        return kwargs

    def get_turbine_kwargs(self) -> t.Dict[str, t.Any]:
        kwargs = {
            "running_list": [["001#", 500, 30100], ["002#", 502, 32220], ["028#", 500, 90000], ["029#", 600, 70000]],
            "distribution_data": {"x": [1, 23, 29, 12, 1],
                                  "labels": ["故障1台", "未检测23台", "正常29台", "告警23", "注意1台"],
                                  "colors": ["#ff0000", "#808080", "#02fa62de", "#ffc000", "#ffff00"]},
            "error_list": [["001#", "主轴", "注意", "主轴后轴承存在间歇性冲击", "持续关注监测数据"],
                           ["001#", "齿轮箱", "故障", "内齿圈局部异常或叶片不平衡，中间轴后轴承保持架故障， 高速级小齿轮局部异常",
                            "关注内齿圈齿面损伤扩展情况，排查中间轴后轴承损伤情况，排查高速级小齿轮损伤情况"],
                           ["001#", "发电机", "注意", "电气故障", "检测发电机三相电流是否平衡"],
                           ["003#", "主轴", "注意", "主轴后轴承存在间歇性冲击", "持续关注监测数据"],
                           ["003#", "齿轮箱", "故障", "内齿圈局部异常或叶片不平衡，中间轴后轴承保持架故障， 高速级小齿轮局部异常",
                            "关注内齿圈齿面损伤扩展情况，排查中间轴后轴承损伤情况，排查高速级小齿轮损伤情况"],
                           ["003#", "发电机", "注意", "电气故障", "检测发电机三相电流是否平衡"]
                           ],
            "all_list": [["001#", "主轴", "注意", "主轴后轴承存在间歇性冲击", "持续关注监测数据", "正常"],
                         ["001#", "齿轮箱", "故障", "内齿圈局部异常或叶片不平衡，中间轴后轴承保持架故障， 高速级小齿轮局部异常",
                          "关注内齿圈齿面损伤扩展情况，排查中间轴后轴承损伤情况，排查高速级小齿轮损伤情况", "正常"],
                         ["001#", "发电机", "注意", "电气故障", "检测发电机三相电流是否平衡", "正常"],
                         ["002#", "主轴", "注意", "主轴后轴承存在间歇性冲击", "持续关注监测数据", "正常"],
                         ["002#", "齿轮箱", "故障", "内齿圈局部异常或叶片不平衡，中间轴后轴承保持架故障， 高速级小齿轮局部异常",
                          "关注内齿圈齿面损伤扩展情况，排查中间轴后轴承损伤情况，排查高速级小齿轮损伤情况", "异常"],
                         ["002#", "发电机", "注意", "电气故障", "检测发电机三相电流是否平衡", "正常"],
                         ["002#", "发电机", "告警", "电气故障", "检测发电机三相电流是否平衡1", "正常"]
                         ],
            "abnormal_turbine_list": [
                {"setting": [["整机厂家", "Vestas", "机型", "V80"], ["额定功率", "2MW", "额定转速", "1700"],
                             ["主轴承厂家", "", "主轴承型号", ""], ["齿轮箱厂家", "力士乐", "齿轮箱型号", "GPV442-i00.54"],
                             ["发电机厂家", "", "发电机型号", ""]],
                 "speed": [["通道", "测点名称", "测量方向", "频带(hz)", "有效值95分位(m/s^2)", "预警门限(m/s^2)"],
                           ["通道1", "主轴轴承径向", "径向", "0.1-10", "0.43", "0.5"],
                           ["通道2", "齿轮箱一级行星架叶片侧轴承径向", "径向", "0.1-10", "1.82", "0.5"],
                           ["通道2", "齿轮箱一级行星架叶片侧轴承径向", "径向", "10-2000", "3.68", "1.5"],
                           ["通道3", "齿轮箱一级内齿圈径向", "径向", "0.1-10", "1.77", "0.5"],
                           ["通道3", "齿轮箱一级内齿圈径向", "径向", "10-2000", "3.68", "1.5"],
                           ["通道4", "发电机机架", "", "10-5000", "3.68", "16"]],
                 "result": [{"name": "主轴", "desc": "正常"},
                            {"name": "齿轮箱", "desc": "太阳轮疑似不平衡或二级大齿轮局部异常，内齿圈局部异常或叶片不平衡， 中间轴后轴承疑似故障。密切关注高速轴故障发展情况"},
                            {"name": "发电机", "desc": "正常"}],
                 "suggest": [{"name": "主轴", "desc": "无"},
                             {"name": "齿轮箱", "desc": "关注二级大齿轮齿面损伤扩展情况 ，内齿圈齿面损伤扩展情况，排查中间轴后轴承故障，密切关注高速轴故障发展情况。"},
                             {"name": "发电机", "desc": "无"}],
                 "analyze": [{"name": "主轴恶化综合指标趋势：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"},
                             {"name": "齿轮箱恶化综合指标趋势：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"},
                             {"name": "发电机恶化综合指标趋势：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"},
                             {"name": "时域指标趋势图：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"},
                             {"name": "频谱特征频率幅值趋势图：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"},
                             {"name": "包络谱特征频率相对幅值趋势：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"},
                             {"name": "边频带频率幅值趋势：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"},
                             ],
                 "signal_analyze": [{"name": "时域分析：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"},
                                    {"name": "频域分析：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"},
                                    {"name": "包络分析：", "desc": "主轴运行平稳，各测点振动指标平稳。", "img_path": "C:\\Users\\EDY\\Desktop\\test.png"}],
                 "turbine_code": "003",
                 }],
            "normal_turbine_list": ["xxxx", "惆怅长岑长惆怅长岑长", "3333333333333333333"],
        }
        return kwargs

    def read_document(self):
        documents = MailMerge(self.base_template_path)
        documents.merge(**self.get_merge_kwargs())
        documents.write(self.tmp_save_path)

        time.sleep(0.1)

        self.document = Document(self.tmp_save_path)

    def write_document(self):
        turbine_data = self.get_turbine_kwargs()
        # self.write_table(self.document, turbine_data["running_list"], 3, merge=False)
        self.write_picture(self.document, self.draw_picture(turbine_data["distribution_data"]), 37)
        self.write_table(self.document, turbine_data["error_list"], 3, merge=True)
        self.write_table(self.document, turbine_data["all_list"], 4, merge=True)
        self.document.add_heading("异常机组详细分析")
        for index, doc in enumerate(turbine_data["abnormal_turbine_list"], start=1):
            self.abnormal_turbine(self.document, 55, index, doc)

        self.document.add_heading("正常机组分析")
        for index, doc in enumerate(turbine_data["normal_turbine_list"], start=1):
            self.normal_turbine(self.document, 57, index, doc)

    def abnormal_turbine(self, document_to, idx, start, data):
        turbine = document_to
        turbine.add_heading(f"陆丰甲东风场{data['turbine_code']}#机组", 2)

        turbine.add_heading(f"机组配置基本信息", 3)

        self.write_table(turbine, data["setting"], merge=False)
        turbine.add_paragraph()

        turbine.add_heading(f"测点配置及加速度有效值", 3)

        self.write_table(turbine, data["speed"])
        turbine.add_paragraph()

        turbine.add_heading(f"诊断结论", 3)
        # turbine.add_run().add_break()

        for item in data['result']:
            r = turbine.add_paragraph(f"{item['name']}：{item['desc']}")
            # r.add_run().add_break()

        turbine.add_heading(f"检维修建议", 3)
        # turbine.add_run().add_break()
        for item in data['suggest']:
            r = turbine.add_paragraph(f"{item['name']}：{item['desc']}")
            # r.add_run().add_break()

        turbine.add_heading(f"历史趋势分析", 3)
        # turbine.add_run().add_break()
        for item in data['analyze']:
            turbine.add_paragraph(f"{item['name']}")
            # turbine.add_run().add_break()
            self.add_picture(turbine.add_paragraph(), item['img_path'])
            turbine.add_paragraph(f"{item['desc']}")
            # turbine.add_run().add_break()

        turbine.add_heading(f"典型单条信号分析", 3)
        # turbine.add_run().add_break()
        for item in data['signal_analyze']:
            turbine.add_paragraph(f"{item['name']}")
            # turbine.add_run().add_break()
            self.add_picture(turbine.add_paragraph(), item['img_path'])
            turbine.add_paragraph(f"{item['desc']}")
            # turbine.add_run().add_break()

    def normal_turbine(self, document_to, idx, start, data):
        r = document_to.add_paragraph(f"{data}")
        r.add_run().add_break()

    def save_document(self):
        if self.kwargs.get("tail_template_path"):
            tail = Document(self.kwargs.get("tail_template_path"))
            self.document = Composer(self.document)
            self.document.append(tail)

        self.document.save(self.tmp_save_path)

    def write_table(self, document_to, table_data, table_idx=None, table_styles='table_text', merge=True):
        """
        插入表格，支持合并列功能，
        :param table_styles: 段落格式
        :param document_to: 写入文档
        :param table_data: 写入数据
        :param table_idx: 表格ID, None则为新建表格
        :param merge: 是否合并表格 当前只会上下相邻合并
        :return:
        """
        table_obj = document_to.add_table(
            1, len(table_data[0]), style="Table Grid"
        ) if table_idx is None else document_to.tables[table_idx]
        for i, row_data in enumerate(table_data):
            new_row = table_obj.add_row()
            new_row.height = Cm(0.6)
            for k, cell in enumerate(new_row.cells):
                cell.paragraphs[0].styles = document_to.styles[table_styles]  # 设置样式
                # cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 上下居中
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 左右居中
                self.Color.get(row_data[k]) and cell._tc.get_or_add_tcPr().append(
                    parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.Color[row_data[k]]}"/>'))  # 设置背景颜色

                # 合并单元格
                if merge and i > 0 and table_data[i][0] == table_data[i - 1][0] and table_data[i][k] == \
                        table_data[i - 1][k]:
                    # 因为表格有头table_obj.cell(x, y) 所以x需要整体+1
                    base_cell = table_obj.cell(i + 1, k)
                    child_cell = table_obj.cell(i, k)
                    base_cell.merge(child_cell)
                else:
                    # 写表格
                    cell.text = str(row_data[k])

        # 假如是新建表格，删除掉一个空白行
        table_idx is None and table_obj._tbl.remove(table_obj.rows[0]._tr)

    def write_picture(self, document_to, picture_path_list, idx) -> None:
        """
        写图片
        :param document_to: 文档
        :param picture_path_list: 图片路径
        :param idx: 写入位置
        :return:
        """
        paragraph = document_to.paragraphs[idx]
        if isinstance(picture_path_list, (str, io.BytesIO)):
            self.add_picture(paragraph, picture_path_list)
            return

        for pic_path in picture_path_list:
            self.add_picture(paragraph, pic_path)

    def add_picture(self, paragraph, pic_path):
        """
        :param paragraph: 段落
        :param pic_path: 图片路径
        :return:
        """
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic = paragraph.add_run().add_picture(pic_path)
        pic.height = Cm(10 * pic.height / pic.width)
        pic.width = Cm(10)
        # paragraph.add_run().add_break()

    def write_text(self, document_to, text, idx=-1, run=False) -> None:
        """
        写文字
        :param document_to: 文档
        :param text: 内容
        :param idx: -1: 往文档最后添加, 其他任意位置前添加
        :param run: 是否增加行
        :return:
        """
        if run:
            r = document_to[idx].add_run(text)
            r.add_break()
            return

        if idx == -1:
            document_to.add_paragraph(text)
        else:
            document_to[idx].insert_paragraph_before(text)

    def delete_paragraph(self, document_to, idx) -> None:
        """
        删除段落
        :param document_to: 文档
        :param idx: 索引下标
        :return:
        """
        p = document_to.paragraphs[idx]._element
        p.getparent().remove(p)
        p._p = p._element = None

    def draw_picture(self, data):
        """
        分布饼图
        :param data:
        :return:
        """
        import matplotlib
        import matplotlib.pyplot as plt
        from pylab import mpl
        # 设置显示中文字体
        mpl.rcParams["font.sans-serif"] = ["SimHei"]
        matplotlib.use('TkAgg')
        my_dpi = 96
        plt.figure(figsize=(480 / my_dpi, 480 / my_dpi), dpi=my_dpi)
        plt.pie(x=data["x"],  # 指定绘图数据
                labels=data["labels"],  # 为饼图添加标签说明
                colors=data["colors"]
                )
        plt.title("机组状态分布图", fontsize=12)

        tmp_buffer = io.BytesIO()
        plt.savefig(tmp_buffer, format='png')
        tmp_buffer.seek(0)
        return tmp_buffer

    def insert_table_after(self, paragraph, table, idx):
        # index = self.document.paragraphs.index(paragraph)  # 获取当前段落的索引
        # self.document.paragraphs[idx-1]._p.addnext(table._tbl)
        paragraph._p.addnext(table._tbl)

    def run(self, *args, **kwargs):
        self.read_document()
        self.write_document()
        self.save_document()


if __name__ == '__main__':
    uid = uuid.uuid1()
    template_docx = f'C:\\Users\\EDY\\Desktop\\振动监测评估报告模版.docx'
    tmp_save_path = f'C:\\Users\\EDY\\Desktop\\振动监测评估报告模版3.docx'
    tail_template_path = f'C:\\Users\\EDY\\Desktop\\振动检测评估报告结尾模版.docx'
    abnormal_template_path = f'C:\\Users\\EDY\\Desktop\\振动监测评估报告异常机组模版.docx'
    report_obj = CmsReport(template_docx, tmp_save_path, tail_template_path=tail_template_path,
                           abnormal_template_path=abnormal_template_path)
    report_obj.run()
